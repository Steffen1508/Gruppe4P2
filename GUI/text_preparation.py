from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    text = "\n\n".join(p.extract_text() or "" for p in PdfReader(str(path)).pages).strip()
    if not text:
        raise ValueError(
            "Fant ingen lesbar tekst i PDF-en. "
            "Den kan være skannet som bilde og kreve OCR."
        )
    return text


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Inkluder tabeller som pipe-separerede rækker
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text for c in row.cells)
            if row_text.strip():
                parts.append(row_text)
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("Fant ingen tekst i DOCX-filen.")
    return text


def _extract_json(path: Path) -> str:
    # Flader alle strenge i JSON-træet ud — tilpas hvis I har et fast skema
    with path.open("r", encoding="utf-8") as f:
        data = json.load(f)

    def walk(obj) -> Iterable[str]:
        if isinstance(obj, str):
            yield obj
        elif isinstance(obj, dict):
            for v in obj.values():
                yield from walk(v)
        elif isinstance(obj, list):
            for v in obj:
                yield from walk(v)

    text = "\n".join(s for s in walk(data) if s.strip())
    if not text:
        raise ValueError("Fant ingen tekststrenger i JSON-filen.")
    return text


EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".json": _extract_json,
}

SUPPORTED_EXTS = set(EXTRACTORS.keys())


def extract_text(path: str | Path) -> str:
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Filen finnes ikke: {path}")
    suffix = path.suffix.lower()
    if suffix not in EXTRACTORS:
        raise ValueError(
            f"Ugyldig filtype: {suffix}. "
            f"Støttede filtyper: {sorted(SUPPORTED_EXTS)}"
        )
    return EXTRACTORS[suffix](path)


def chunk_with_offsets(
    text: str,
    tokenizer,
    max_tokens: int = 512,
    overlap: int = 50,
) -> list[tuple[str, int]]:
    if overlap >= max_tokens:
        raise ValueError("overlap skal være mindre end max_tokens")
    if not text or not text.strip():
        return []

    effective_max = max_tokens - 2  # reservér plads til [CLS] og [SEP]
    step = effective_max - overlap  # hvor langt vi rykker mellem chunks

    encoding = tokenizer(
        text,
        add_special_tokens=False,
        return_offsets_mapping=True,
        truncation=False,
    )
    token_ids = encoding["input_ids"]
    offsets = encoding["offset_mapping"]

    # Kort nok til at returnere direkte uden chunking
    if len(token_ids) <= effective_max:
        return [(text, 0)]

    chunks = []
    start_tok = 0
    while start_tok < len(token_ids):
        end_tok = min(start_tok + effective_max, len(token_ids))
        char_start = offsets[start_tok][0]
        char_end = offsets[end_tok - 1][1]
        chunk = text[char_start:char_end]
        if chunk.strip():
            chunks.append((chunk, char_start))
        if end_tok >= len(token_ids):
            break
        start_tok += step

    return chunks


def count_tokens(text: str, tokenizer) -> int:
    if not text:
        return 0
    return len(tokenizer.encode(text, add_special_tokens=False))